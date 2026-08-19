#!/usr/bin/env python3
"""Fix Meridian dataset quality: formulas, remove SYNTHETIC banners, diversify notes/dates."""
from __future__ import annotations

import random
import zipfile
from datetime import timedelta
from pathlib import Path

import pandas as pd
from docx import Document
from openpyxl import load_workbook
from openpyxl.styles import Alignment, Font
from openpyxl.utils import get_column_letter
from pypdf import PdfReader
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import Image as RLImage
from reportlab.platypus import PageBreak, Paragraph, Preformatted, SimpleDocTemplate, Spacer

ROOT = Path(__file__).resolve().parent
ASOF = pd.Timestamp("2026-08-09")


def strip_inventory_ledger_banners() -> None:
    for path, title, subtitle in [
        (
            ROOT / "it_asset_inventory.xlsx",
            "Yanou & Partners IT Services - IT Asset Inventory",
            "Operational snapshot | As of 2026-08-09",
        ),
        (
            ROOT / "fixed_asset_ledger.xlsx",
            "Yanou & Partners IT Services - Fixed Asset Ledger",
            "Accounting extract | As of 2026-08-09",
        ),
    ]:
        wb = load_workbook(path)
        ws = wb.active
        ws["A1"] = title
        ws["A2"] = subtitle
        wb.save(path)
        print(f"updated banners: {path.name}")


def irregularize_purchase_dates() -> None:
    """Break mechanical MM-DD == sequential pattern without changing asset tags."""
    rng = random.Random(20260809)
    path = ROOT / "it_asset_inventory.xlsx"
    wb = load_workbook(path)
    ws = wb.active
    # header row 4
    headers = {ws.cell(4, c).value: c for c in range(1, ws.max_column + 1)}
    pd_col = headers["Purchase Date"]
    war_col = headers["Warranty Expiration"]
    cat_col = headers["Category"]
    for r in range(5, ws.max_row + 1):
        tag = ws.cell(r, 1).value
        if not tag:
            continue
        purchase = ws.cell(r, pd_col).value
        warranty = ws.cell(r, war_col).value
        if purchase is None:
            continue
        # jitter day/month irregularly; keep year
        delta_days = rng.choice([-19, -14, -11, -7, -3, 2, 5, 9, 13, 17, 21, 26])
        new_p = pd.Timestamp(purchase) + timedelta(days=int(delta_days))
        # keep warranty relative span if present
        if warranty is not None:
            span = pd.Timestamp(warranty) - pd.Timestamp(purchase)
            new_w = new_p + span
            # small extra noise on warranty end
            new_w = new_w + timedelta(days=rng.choice([-5, -2, 0, 1, 4, 8]))
            ws.cell(r, war_col).value = new_w.to_pydatetime().replace(hour=12, minute=0, second=0, microsecond=0)
        ws.cell(r, pd_col).value = new_p.to_pydatetime().replace(hour=12, minute=0, second=0, microsecond=0)
    wb.save(path)
    print("irregularized purchase/warranty dates in it_asset_inventory.xlsx")


def diversify_regional_notes() -> None:
    path = ROOT / "regional_IT_notes.docx"
    doc = Document(path)
    # Replace synthetic banner paragraph
    for p in doc.paragraphs:
        t = p.text.strip()
        if t.startswith("SYNTHETIC"):
            p.text = (
                "Working notes only. Treat as investigative leads. "
                "They do not prove custody, return, approval, receipt, loss, or disposal."
            )
        elif "SYNTHETIC" in t:
            p.text = t.replace("SYNTHETIC, NON-AUTHORITATIVE EVIDENCE. ", "").replace("SYNTHETIC", "").strip()

    variants = {
        "MD-00089": "During the Denver exit walkthrough, someone mentioned a locked cage on floor 3 might still hold this laptop. No cycle-count tag, transfer, or dock receipt was filed.",
        "MD-00090": "Chicago facilities said spare monitors were stacked in the old MDF after the move. This unit was named verbally; nothing was scanned into the register.",
        "MD-00091": "A contractor thought MD-00091 went to temporary overflow near the loading dock. Badge logs were not pulled and no transfer ID exists.",
        "MD-00092": "Floor lead recalled boxing network gear with Chicago surplus. Serial was not photographed; claim is unverified.",
        "MD-00093": "Night-shift note: device may have stayed in a manager's office that was later sealed. No inventory ticket attached.",
        "MD-00094": "IT float pool chat referenced this asset as 'still at closed site.' Chat is not an approval or receiving record.",
        "MD-00095": "Vendor teardown crew said they saw a similar model on a pallet marked scrap hold. Asset tag was not confirmed on camera.",
        "MD-00096": "Helpdesk comment claims the user dropped it with building security before last day. Security has no intake log for the tag.",
        "MD-00097": "Regional spare inventory spreadsheet (offline) listed a placeholder row for this tag with location TBD. Spreadsheet is not a system of record.",
        "MD-00098": "Move captain's email guessed the device rode with furniture to storage. Furniture BOL does not list IT assets.",
        "MD-00099": "An intern's handwritten punch list includes this tag under 'find later.' No follow-up work order was opened.",
        "MD-00100": "Closing checklist has a sticky note 'missing — check cage.' Sticky note is not receiving evidence.",
        "MD-00114": "Recycler pickup was reportedly included with the July pallet, but the certificate bundle did not list this asset tag.",
        "MD-00118": "Technician believes the device was wiped and recycled. No certificate or approval was located.",
        "MD-00058": "Asset appears to have been moved to Regional IT stock. The transfer record contains no approval ID.",
        "MD-00082": "Return package was delivered according to the technician, but the shipment serial does not match the asset register and the dock exception photo is only a lead.",
    }

    table = doc.tables[0]
    for row in table.rows[1:]:
        tag = row.cells[1].text.strip()
        if tag in variants:
            row.cells[3].text = variants[tag]
            row.cells[4].text = "Lead only — unverified"
    doc.save(path)
    print("diversified regional_IT_notes.docx")


def rebuild_policy_pdf() -> None:
    """Rebuild policy PDF without SYNTHETIC disclaimer; keep control-matrix figure if present."""
    reader = PdfReader(str(ROOT / "IT_asset_management_policy.pdf"))
    img_path = ROOT / "ITAM_control_matrix.png"
    # extract embedded image if needed
    for page in reader.pages:
        try:
            for i, img in enumerate(page.images):
                tmp = ROOT / f"_policy_img_{i}.png"
                tmp.write_bytes(img.data)
                img_path = tmp
        except Exception:
            pass

    styles = getSampleStyleSheet()
    title = ParagraphStyle("T", parent=styles["Heading1"], fontSize=16, spaceAfter=8)
    h = ParagraphStyle("H", parent=styles["Heading2"], fontSize=12, spaceBefore=10, spaceAfter=4)
    body = ParagraphStyle("B", parent=styles["Normal"], fontSize=9, leading=12, spaceAfter=4)
    small = ParagraphStyle("S", parent=styles["Normal"], fontSize=8, leading=10, textColor=colors.HexColor("#333333"))

    story = []
    story.append(Paragraph("IT Asset Management Policy", title))
    story.append(
        Paragraph(
            "Yanou & Partners IT Services | Policy ITAM-001 | Version 3.2 | Effective 2026-01-01",
            small,
        )
    )
    story.append(Spacer(1, 0.15 * inch))

    sections = [
        (
            "1. Purpose and scope",
            "This policy establishes accountable custody, movement, return, loss, financial reconciliation, sanitization, "
            "and disposal controls for company-owned laptops, monitors, mobile devices, and network equipment at offices "
            "and remote locations.",
        ),
        (
            "2. System of record",
            "The IT asset register is the operational system of record for assignment and location. The fixed-asset ledger "
            "is the financial system of record for capitalized cost and net book value. Neither file may be updated from "
            "verbal claims alone.",
        ),
        (
            "3. Assignment and transfer",
            "Assets move only with a transfer ID, approved approver ID, and receiving acknowledgment. Blank approval IDs "
            "are control exceptions. Closed-office locations are not valid verified locations after consolidation.",
        ),
        (
            "4. Offboarding and returns",
            "Separated employees must return assigned assets within 10 business days. A carrier label is not proof of return. "
            "Return completion requires carrier acceptance, delivery confirmation, and a receiving scan that matches asset "
            "tag and serial number.",
        ),
        (
            "5. Receiving exceptions",
            "Shipment serial mismatches, damaged parcels, and dock exception photos are investigative leads. An exception "
            "image does not clear custody or mark an asset Available.",
        ),
        (
            "6. Disposal and media sanitization",
            "Retirement requires an approved disposal path and a verified disposal certificate matching tag/serial. "
            "Sanitization should align with NIST SP 800-88 Rev. 2 media sanitization guidance before recycler release.",
        ),
        (
            "7. Financial reconciliation",
            "Operational inventory must be reconciled to the fixed-asset ledger each quarter. Every remaining difference "
            "(missing ledger row, cost-basis mismatch, or status disagreement) must be explained for Finance. Useful life "
            "and depreciation references follow IRS Publication 946 concepts as implemented in company useful-life tables "
            "(typically 3-year SL for laptops/mobile; 5-year SL for monitors/network).",
        ),
        (
            "8. Access control and auditability",
            "Control design references NIST SP 800-53 Rev. 5 families for accountability and audit records. Evidence must "
            "remain attributable to a system transaction ID.",
        ),
        (
            "9. Certification",
            "The IT Operations Manager certifies the quarterly inventory only when Critical custody and financial blockers "
            "are cleared or formally accepted by Finance and Internal Audit. HR validates employment status used in custody "
            "decisions.",
        ),
        (
            "10. Evidence and retention",
            "Evidence must identify the asset tag, serial number where applicable, transaction ID, date, and source system. "
            "Technician notes are leads only and are retained with the reconciliation package for audit follow-up.",
        ),
    ]
    for heading, text in sections:
        story.append(Paragraph(heading, h))
        story.append(Paragraph(text, body))

    story.append(PageBreak())
    story.append(Paragraph("Appendix A — Control Matrix (figure)", h))
    story.append(
        Paragraph(
            "The following figure is part of the policy package and must be read with the narrative controls above.",
            body,
        )
    )
    if img_path.exists():
        story.append(Spacer(1, 0.1 * inch))
        story.append(RLImage(str(img_path), width=6.5 * inch, height=3.6 * inch))
    story.append(Spacer(1, 0.2 * inch))
    story.append(Paragraph("Appendix B — Public Benchmark Sources", h))
    story.append(
        Preformatted(
            "NIST SP 800-53 Rev. 5 — Security and Privacy Controls\n"
            "NIST SP 800-88 Rev. 2 — Media Sanitization\n"
            "IRS Publication 946 — How To Depreciate Property",
            small,
        )
    )

    out = ROOT / "IT_asset_management_policy.pdf"
    SimpleDocTemplate(str(out), pagesize=letter, title="ITAM-001").build(story)
    for tmp in ROOT.glob("_policy_img_*.png"):
        tmp.unlink(missing_ok=True)
    print("rebuilt IT_asset_management_policy.pdf without SYNTHETIC disclaimer")


def add_source_ledger_sheet(wb) -> tuple[int, int]:
    """Embed fixed-asset ledger values for XLOOKUP formulas. Returns (first_data_row, last_data_row)."""
    if "Source Ledger" in wb.sheetnames:
        del wb["Source Ledger"]
    src = load_workbook(ROOT / "fixed_asset_ledger.xlsx", data_only=True).active
    ws = wb.create_sheet("Source Ledger", 0)
    # copy used grid
    for r in range(1, src.max_row + 1):
        for c in range(1, src.max_column + 1):
            ws.cell(r, c).value = src.cell(r, c).value
    # data starts row 5 in source files
    return 5, src.max_row


def formulize_workbook() -> None:
    path = ROOT / "Meridian_IT_Asset_Reconciliation.xlsx"
    wb = load_workbook(path)
    led_first, led_last = add_source_ledger_sheet(wb)

    cr = wb["Corrected Register"]
    er = wb["Exception Register"]
    dash = wb["Dashboard"]
    lr = wb["Ledger Reconciliation"]
    cc = wb["Custody Chain"]

    cr_first, cr_last = 5, 136
    er_first, er_last = 5, 167

    # --- Corrected Register: NBV + Ledger ID via XLOOKUP to Source Ledger ---
    # Source Ledger: A Ledger Asset ID, B Asset Tag, E Acquisition Cost, G Net Book Value, H Ledger Status
    for r in range(cr_first, cr_last + 1):
        if not cr.cell(r, 1).value:
            continue
        # Remaining Book Value (col 14)
        cr.cell(r, 14).value = (
            f'=IFERROR(XLOOKUP(A{r},\'Source Ledger\'!$B${led_first}:$B${led_last},'
            f'\'Source Ledger\'!$G${led_first}:$G${led_last}),"")'
        )
        # Ledger Asset ID (col 15)
        cr.cell(r, 15).value = (
            f'=IFERROR(XLOOKUP(A{r},\'Source Ledger\'!$B${led_first}:$B${led_last},'
            f'\'Source Ledger\'!$A${led_first}:$A${led_last}),"")'
        )

    # --- Exception Register: helper cost cols + formula exposure ---
    # Insert formulas in Financial Exposure (col G). Use:
    # register cost XLOOKUP, ledger cost XLOOKUP inline.
    for r in range(er_first, er_last + 1):
        if not er.cell(r, 1).value:
            continue
        # Financial Exposure
        # Inventory-to-Ledger Difference: if ledger cost found -> ABS(ledger-register); else register acquisition cost
        # Else: remaining book value from corrected register
        er.cell(r, 7).value = (
            f'=IF(B{r}="Inventory-to-Ledger Difference",'
            f'IF(IFERROR(XLOOKUP(D{r},\'Source Ledger\'!$B${led_first}:$B${led_last},'
            f'\'Source Ledger\'!$E${led_first}:$E${led_last}),"")="",'
            f'IFERROR(XLOOKUP(D{r},\'Corrected Register\'!$A${cr_first}:$A${cr_last},'
            f'\'Corrected Register\'!$M${cr_first}:$M${cr_last}),0),'
            f'ABS(IFERROR(XLOOKUP(D{r},\'Source Ledger\'!$B${led_first}:$B${led_last},'
            f'\'Source Ledger\'!$E${led_first}:$E${led_last}),0)'
            f'-IFERROR(XLOOKUP(D{r},\'Corrected Register\'!$A${cr_first}:$A${cr_last},'
            f'\'Corrected Register\'!$M${cr_first}:$M${cr_last}),0))),'
            f'IFERROR(XLOOKUP(D{r},\'Corrected Register\'!$A${cr_first}:$A${cr_last},'
            f'\'Corrected Register\'!$N${cr_first}:$N${cr_last}),0))'
        )

    # --- Dashboard formulas ---
    dash["A2"] = (
        "Prepared for IT Operations Manager, Finance Controller, HR Operations Lead, and Internal Audit | "
        "As of 2026-08-09 | Policy figure and dock image used as leads only"
    )

    def cr_range(col: str) -> str:
        return f"'Corrected Register'!${col}${cr_first}:${col}${cr_last}"

    def er_range(col: str) -> str:
        return f"'Exception Register'!${col}${er_first}:${col}${er_last}"

    # KPIs in column B
    dash["B5"] = f"=COUNTA({cr_range('A')})"
    dash["B6"] = f"=SUM({cr_range('M')})"
    dash["B7"] = f"=SUM({cr_range('N')})"
    dash["B8"] = f'=COUNTIF({er_range("K")},"Open")'
    dash["B9"] = f'=COUNTIF({er_range("C")},"Critical")'
    dash["B10"] = f'=COUNTIF({cr_range("J")},"Missing")'
    dash["B11"] = f'=COUNTIF({cr_range("J")},"Return Overdue")'
    dash["B12"] = f'=COUNTIF({cr_range("J")},"Disposed")'
    # unique custody assets
    dash["B13"] = f'=COUNTA(UNIQUE(FILTER(\'Custody Chain\'!A{5}:A{cc.max_row},\'Custody Chain\'!A{5}:A{cc.max_row}<>"")))'

    # By verified status block rows 19-26: A=status label, B=count, C=book value
    for r in range(19, 27):
        status = dash.cell(r, 1).value
        if status:
            dash.cell(r, 2).value = f'=COUNTIF({cr_range("J")},A{r})'
            dash.cell(r, 3).value = f'=SUMIF({cr_range("J")},A{r},{cr_range("N")})'

    # By category rows 19-22 cols E-G
    for r in range(19, 23):
        cat = dash.cell(r, 5).value
        if cat:
            dash.cell(r, 6).value = f'=COUNTIF({cr_range("C")},E{r})'
            dash.cell(r, 7).value = f'=SUMIF({cr_range("C")},E{r},{cr_range("N")})'

    # By location rows 32-43
    for r in range(32, 44):
        loc = dash.cell(r, 1).value
        if loc:
            dash.cell(r, 2).value = f'=COUNTIF({cr_range("I")},A{r})'
            dash.cell(r, 3).value = f'=SUMIF({cr_range("I")},A{r},{cr_range("N")})'

    # Exceptions by type rows 32-40 cols E-G
    for r in range(32, 41):
        typ = dash.cell(r, 5).value
        if typ:
            dash.cell(r, 6).value = f'=COUNTIF({er_range("B")},E{r})'
            dash.cell(r, 7).value = f'=SUMIF({er_range("B")},E{r},{er_range("G")})'

    # --- Ledger Reconciliation formulas ---
    lr["B5"] = f"=SUM({cr_range('M')})"
    lr["B6"] = f"=SUM('Source Ledger'!$E${led_first}:$E${led_last})"
    lr["B7"] = "=B5-B6"
    lr["B8"] = f"=SUM({cr_range('N')})"
    lr["B9"] = (
        f'=SUMIF(\'Source Ledger\'!$H${led_first}:$H${led_last},"Active",'
        f'\'Source Ledger\'!$G${led_first}:$G${led_last})'
    )
    lr["B10"] = f"=SUM('Source Ledger'!$G${led_first}:$G${led_last})"
    # matching / mismatch / missing counts via formulas on asset-level table where possible
    # Asset-level table rows 24-35:
    # C register cost, D ledger cost, E difference, F register NBV, G ledger NBV
    for r in range(24, 36):
        tag = lr.cell(r, 1).value
        if not tag:
            continue
        lr.cell(r, 3).value = (
            f'=IFERROR(XLOOKUP(A{r},\'Corrected Register\'!$A${cr_first}:$A${cr_last},'
            f'\'Corrected Register\'!$M${cr_first}:$M${cr_last}),"")'
        )
        lr.cell(r, 4).value = (
            f'=IFERROR(XLOOKUP(A{r},\'Source Ledger\'!$B${led_first}:$B${led_last},'
            f'\'Source Ledger\'!$E${led_first}:$E${led_last}),"")'
        )
        lr.cell(r, 5).value = f'=IF(OR(C{r}="",D{r}=""),"",D{r}-C{r})'
        lr.cell(r, 6).value = (
            f'=IFERROR(XLOOKUP(A{r},\'Corrected Register\'!$A${cr_first}:$A${cr_last},'
            f'\'Corrected Register\'!$N${cr_first}:$N${cr_last}),"")'
        )
        lr.cell(r, 7).value = (
            f'=IFERROR(XLOOKUP(A{r},\'Source Ledger\'!$B${led_first}:$B${led_last},'
            f'\'Source Ledger\'!$G${led_first}:$G${led_last}),"")'
        )

    # Matching / mismatch / missing counts (issue text lives in column B)
    lr["B12"] = '=COUNTIF(B24:B35,"Acquisition cost mismatch")'
    lr["B13"] = '=COUNTIF(B24:B35,"On inventory, missing from ledger")'
    lr["B11"] = f"=COUNTA({cr_range('A')})-B12-B13"

    # Update explanation line 17-18 to use formula-driven combined amounts where easy
    lr["A17"] = (
        '=CONCATENATE("1) Three assets (MD-00130, MD-00131, MD-00132) appear on the operational register and '
        'PO-2023-0022 but have no fixed-asset ledger row. Combined register acquisition cost = $",'
        'TEXT(C33+C34+C35,"0.0"),".")'
    )
    lr["A18"] = (
        '=CONCATENATE("2) Seven assets have a systematic $175 ledger-over-register acquisition cost variance '
        '(MD-00017, MD-00034, MD-00051, MD-00068, MD-00085, MD-00102, MD-00119). Combined ledger cost excess = $",'
        'TEXT(E24+E25+E26+E27+E28+E29+E32,"#,##0"),".")'
    )

    wb.save(path)
    print(f"formulized {path.name}")


def census_formulas() -> None:
    wb = load_workbook(ROOT / "Meridian_IT_Asset_Reconciliation.xlsx")
    print("\n=== Formula census ===")
    for name in wb.sheetnames:
        ws = wb[name]
        typed = 0
        formulas = 0
        for row in ws.iter_rows(max_row=ws.max_row, max_col=ws.max_column):
            for cell in row:
                v = cell.value
                if isinstance(v, (int, float)) and not isinstance(v, bool):
                    typed += 1
                elif isinstance(v, str) and v.startswith("="):
                    formulas += 1
        total = typed + formulas
        share = (typed / total * 100) if total else 0
        print(f"{name}: typed={typed} formulas={formulas} typed_share={share:.1f}%")


def rebuild_zips() -> None:
    inputs = [
        "it_asset_inventory.xlsx",
        "hr_employee_status.csv",
        "equipment_transfer_log.csv",
        "service_desk_offboarding.csv",
        "device_return_shipments.csv",
        "hardware_purchase_orders.csv",
        "fixed_asset_ledger.xlsx",
        "asset_disposal_records.csv",
        "regional_IT_notes.docx",
        "IT_asset_management_policy.pdf",
        "ITAM_control_matrix.png",
        "receiving_exception_scan_1ZMD00000082.png",
    ]
    in_zip = ROOT / "Meridian_IT_Asset_Inputs.zip"
    with zipfile.ZipFile(in_zip, "w", compression=zipfile.ZIP_DEFLATED) as zf:
        for name in inputs:
            zf.write(ROOT / name, arcname=name)
    out_zip = ROOT / "Meridian_IT_Asset_Reconciliation.zip"
    with zipfile.ZipFile(out_zip, "w", compression=zipfile.ZIP_DEFLATED) as zf:
        zf.write(ROOT / "Meridian_IT_Asset_Reconciliation.xlsx", arcname="Meridian_IT_Asset_Reconciliation.xlsx")
    print(f"rebuilt {in_zip.name} ({len(inputs)} files) and {out_zip.name}")


def main() -> None:
    strip_inventory_ledger_banners()
    irregularize_purchase_dates()
    diversify_regional_notes()
    rebuild_policy_pdf()
    formulize_workbook()
    census_formulas()
    rebuild_zips()


if __name__ == "__main__":
    main()
